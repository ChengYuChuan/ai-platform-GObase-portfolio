"""Document loader for various file formats."""

import io
from typing import Any

from langchain_core.documents import Document

from src.core import get_logger
from src.core.config import get_settings
from src.rag.ingestion.chunker import TextChunker

logger = get_logger(__name__)


class DocumentLoader:
    """Load and process documents from various formats."""

    def __init__(self) -> None:
        """Initialize the document loader."""
        self.settings = get_settings()
        self.chunker = TextChunker(
            chunk_size=self.settings.chunk_size,
            chunk_overlap=self.settings.chunk_overlap,
        )

    async def load_and_split(
        self,
        content: bytes,
        filename: str,
        content_type: str,
    ) -> list[Document]:
        """Load document content and split into chunks.

        Args:
            content: Raw file content as bytes.
            filename: Original filename.
            content_type: MIME content type.

        Returns:
            List of Document chunks ready for embedding.
        """
        extension = filename.lower().split(".")[-1] if "." in filename else ""

        logger.info(
            "Loading document",
            filename=filename,
            content_type=content_type,
            extension=extension,
            size_bytes=len(content),
        )

        # Extract text based on file type
        if extension == "pdf" or content_type == "application/pdf":
            text = await self._load_pdf(content)
        elif (
            extension == "docx"
            or content_type
            == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ):
            text = await self._load_docx(content)
        elif extension == "md" or content_type == "text/markdown":
            text = await self._load_markdown(content)
        elif extension == "txt" or content_type == "text/plain":
            text = await self._load_text(content)
        else:
            # Try to decode as text
            text = await self._load_text(content)

        # Split into chunks
        chunks = self.chunker.split_text(text)

        # Convert to LangChain Document format
        documents = [
            Document(
                page_content=chunk,
                metadata={
                    "filename": filename,
                    "content_type": content_type,
                    "chunk_index": i,
                    "total_chunks": len(chunks),
                },
            )
            for i, chunk in enumerate(chunks)
        ]

        logger.info(
            "Document processed",
            filename=filename,
            text_length=len(text),
            chunks_count=len(documents),
        )

        return documents

    async def _load_pdf(self, content: bytes) -> str:
        """Extract text from PDF content."""
        try:
            import pypdf

            pdf_file = io.BytesIO(content)
            reader = pypdf.PdfReader(pdf_file)

            text_parts = []
            for page_num, page in enumerate(reader.pages):
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(f"[Page {page_num + 1}]\n{page_text}")

            return "\n\n".join(text_parts)

        except ImportError:
            logger.error("pypdf not installed, cannot process PDF files")
            raise ValueError("PDF processing requires pypdf library")
        except Exception as e:
            logger.error("Failed to process PDF", error=str(e))
            raise ValueError(f"Failed to process PDF: {str(e)}")

    async def _load_docx(self, content: bytes) -> str:
        """Extract text from DOCX content."""
        try:
            import docx

            docx_file = io.BytesIO(content)
            doc = docx.Document(docx_file)

            text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(
                        cell.text.strip() for cell in row.cells if cell.text.strip()
                    )
                    if row_text:
                        text_parts.append(row_text)

            return "\n\n".join(text_parts)

        except ImportError:
            logger.error("python-docx not installed, cannot process DOCX files")
            raise ValueError("DOCX processing requires python-docx library")
        except Exception as e:
            logger.error("Failed to process DOCX", error=str(e))
            raise ValueError(f"Failed to process DOCX: {str(e)}")

    async def _load_markdown(self, content: bytes) -> str:
        """Extract text from Markdown content.

        For RAG, we keep the markdown structure as it provides context.
        """
        try:
            text = content.decode("utf-8")
            return text
        except UnicodeDecodeError:
            # Try other encodings
            for encoding in ["latin-1", "cp1252", "iso-8859-1"]:
                try:
                    return content.decode(encoding)
                except UnicodeDecodeError:
                    continue
            raise ValueError("Could not decode markdown file with known encodings")

    async def _load_text(self, content: bytes) -> str:
        """Extract text from plain text content."""
        try:
            return content.decode("utf-8")
        except UnicodeDecodeError:
            # Try other encodings
            for encoding in ["latin-1", "cp1252", "iso-8859-1"]:
                try:
                    return content.decode(encoding)
                except UnicodeDecodeError:
                    continue
            raise ValueError("Could not decode text file with known encodings")


class MultiDocumentLoader:
    """Load multiple documents in batch."""

    def __init__(self) -> None:
        """Initialize the multi-document loader."""
        self.loader = DocumentLoader()

    async def load_directory(
        self,
        directory_path: str,
        extensions: list[str] | None = None,
    ) -> list[Document]:
        """Load all documents from a directory.

        Args:
            directory_path: Path to directory containing documents.
            extensions: List of file extensions to include (e.g., ['.pdf', '.md']).

        Returns:
            List of all Document chunks from all files.
        """
        import os

        if extensions is None:
            extensions = [".pdf", ".docx", ".md", ".txt"]

        all_documents: list[Document] = []

        for root, _dirs, files in os.walk(directory_path):
            for filename in files:
                ext = os.path.splitext(filename)[1].lower()
                if ext in extensions:
                    file_path = os.path.join(root, filename)

                    try:
                        with open(file_path, "rb") as f:
                            content = f.read()

                        content_type = self._get_content_type(ext)
                        docs = await self.loader.load_and_split(
                            content=content,
                            filename=filename,
                            content_type=content_type,
                        )

                        # Add file path to metadata
                        for doc in docs:
                            doc.metadata["file_path"] = file_path

                        all_documents.extend(docs)

                    except Exception as e:
                        logger.error(
                            "Failed to load file",
                            file_path=file_path,
                            error=str(e),
                        )

        logger.info(
            "Loaded directory",
            directory=directory_path,
            total_documents=len(all_documents),
        )

        return all_documents

    def _get_content_type(self, extension: str) -> str:
        """Get MIME content type from file extension."""
        content_types = {
            ".pdf": "application/pdf",
            ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            ".md": "text/markdown",
            ".txt": "text/plain",
        }
        return content_types.get(extension, "application/octet-stream")
